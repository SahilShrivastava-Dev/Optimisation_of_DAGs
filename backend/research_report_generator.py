"""
Research Report Generator for DAG Optimization
Generates professional DOCX reports formatted like academic papers
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from typing import Dict, List, Any
import io

class ResearchReportGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles
        
        # Title style
        if 'Custom Title' not in styles:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 51, 102)
    
    def generate_report(self, optimization_data: Dict[str, Any]) -> io.BytesIO:
        """Generate complete research report"""
        
        # Extract data
        original = optimization_data['original']
        optimized = optimization_data['optimized']
        orig_metrics = original['metrics']
        opt_metrics = optimized['metrics']
        timestamp = optimization_data.get('timestamp', datetime.now().isoformat())
        
        # Build report sections
        self._add_title_page(orig_metrics, opt_metrics, timestamp)
        self._add_abstract(orig_metrics, opt_metrics)
        self._add_introduction()
        self._add_methodology()
        self._add_original_graph_analysis(orig_metrics, original['edges'])
        self._add_optimization_process()
        self._add_results_analysis(orig_metrics, opt_metrics)
        self._add_detailed_metrics_comparison(orig_metrics, opt_metrics)
        self._add_efficiency_analysis(orig_metrics, opt_metrics)
        self._add_critical_path_analysis(orig_metrics, opt_metrics)
        self._add_conclusions(orig_metrics, opt_metrics)
        self._add_references()
        
        # Save to BytesIO
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _add_title_page(self, orig_metrics: Dict, opt_metrics: Dict, timestamp: str):
        """Add title page"""
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('Directed Acyclic Graph Optimization:\nA Comprehensive Analysis')
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        self.doc.add_paragraph()  # Spacing
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Research Report on Graph Optimization Techniques')
        run.font.size = Pt(14)
        run.font.italic = True
        
        self.doc.add_paragraph()
        
        # Metadata
        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_str = datetime.fromisoformat(timestamp.replace('Z', '+00:00')).strftime('%B %d, %Y')
        meta.add_run(f'Generated: {date_str}\n')
        meta.add_run(f'DAG Optimizer v3.0\n\n')
        
        # Key stats box
        stats = self.doc.add_paragraph()
        stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stats_text = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
   Original Graph: {orig_metrics['num_nodes']} nodes, {orig_metrics['num_edges']} edges
   Optimized Graph: {opt_metrics['num_nodes']} nodes, {opt_metrics['num_edges']} edges
   Reduction: {((orig_metrics['num_edges'] - opt_metrics['num_edges']) / orig_metrics['num_edges'] * 100):.1f}% edges removed
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        """
        run = stats.add_run(stats_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
        self.doc.add_page_break()
    
    def _add_abstract(self, orig_metrics: Dict, opt_metrics: Dict):
        """Add abstract section"""
        self.doc.add_heading('Abstract', level=1)
        
        edge_reduction = ((orig_metrics['num_edges'] - opt_metrics['num_edges']) / orig_metrics['num_edges'] * 100)
        efficiency_gain = ((opt_metrics['efficiency_score'] - orig_metrics['efficiency_score']) * 100)
        redundancy_reduction = ((orig_metrics['redundancy_ratio'] - opt_metrics['redundancy_ratio']) * 100)
        
        abstract_text = f"""This report presents a comprehensive analysis of Directed Acyclic Graph (DAG) optimization techniques applied to a graph with {orig_metrics['num_nodes']} nodes and {orig_metrics['num_edges']} edges. Through the application of transitive reduction and node equivalence merging algorithms, we achieved a {edge_reduction:.1f}% reduction in edge count while preserving graph semantics. The optimization resulted in a {efficiency_gain:.1f}% improvement in overall efficiency score and a {redundancy_reduction:.1f}% reduction in redundancy ratio. This study demonstrates the practical application of graph theory algorithms in reducing computational complexity while maintaining structural integrity."""
        
        p = self.doc.add_paragraph(abstract_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Keywords
        self.doc.add_paragraph()
        keywords = self.doc.add_paragraph()
        keywords.add_run('Keywords: ').bold = True
        keywords.add_run('Directed Acyclic Graph, Graph Optimization, Transitive Reduction, Node Merging, Computational Complexity, Network Analysis')
        
        self.doc.add_page_break()
    
    def _add_introduction(self):
        """Add introduction section"""
        self.doc.add_heading('1. Introduction', level=1)
        
        intro_paragraphs = [
            "Directed Acyclic Graphs (DAGs) are fundamental data structures widely used in various domains including task scheduling, dependency resolution, version control systems, and data processing pipelines. However, real-world DAGs often contain redundant edges and equivalent nodes that increase computational complexity without adding semantic value.",
            
            "This research report analyzes the application of two primary optimization techniques: (1) Transitive Reduction, which removes redundant edges while preserving reachability, and (2) Node Equivalence Merging, which consolidates nodes with identical predecessor and successor sets. These techniques are crucial for improving performance in large-scale graph processing applications.",
            
            "The objectives of this analysis are threefold: first, to quantify the structural improvements achieved through optimization; second, to evaluate the impact on graph complexity metrics; and third, to identify critical paths and bottlenecks within the optimized structure."
        ]
        
        for text in intro_paragraphs:
            p = self.doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self.doc.add_paragraph()  # Spacing
    
    def _add_methodology(self):
        """Add methodology section"""
        self.doc.add_heading('2. Methodology', level=1)
        
        self.doc.add_heading('2.1 Optimization Algorithms', level=2)
        
        # Transitive Reduction
        self.doc.add_paragraph().add_run('Transitive Reduction').bold = True
        tr_text = "Transitive reduction removes redundant edges from a DAG while preserving its transitive closure. For a DAG G = (V, E), the transitive reduction G' = (V, E') satisfies: (1) E' ⊆ E, (2) the transitive closure of G' equals the transitive closure of G, and (3) E' is minimal with respect to property (2)."
        p = self.doc.add_paragraph(tr_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
        
        # Node Merging
        self.doc.add_paragraph().add_run('Node Equivalence Merging').bold = True
        nm_text = "Two nodes u and v are considered equivalent if they have identical predecessor and successor sets: pred(u) = pred(v) and succ(u) = succ(v). Merging equivalent nodes reduces graph size without affecting connectivity or reachability properties."
        p = self.doc.add_paragraph(nm_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
        
        self.doc.add_heading('2.2 Metrics and Evaluation', level=2)
        
        metrics_text = "We employ a comprehensive set of 20+ metrics to evaluate graph quality, including structural metrics (node/edge count, density), complexity metrics (cyclomatic complexity, topological complexity), efficiency metrics (redundancy ratio, compactness score), and centrality measures (betweenness centrality for bottleneck identification)."
        p = self.doc.add_paragraph(metrics_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_page_break()
    
    def _add_original_graph_analysis(self, orig_metrics: Dict, edges: List[Dict]):
        """Add original graph analysis"""
        self.doc.add_heading('3. Original Graph Analysis', level=1)
        
        analysis_text = f"The input graph consists of {orig_metrics['num_nodes']} nodes and {orig_metrics['num_edges']} edges, forming a directed acyclic structure with {orig_metrics['num_leaf_nodes']} leaf nodes. The graph exhibits a topological complexity of {orig_metrics['topological_complexity']} levels and a maximum path length (diameter) of {orig_metrics['diameter']}."
        
        p = self.doc.add_paragraph(analysis_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
        
        # Structural characteristics
        self.doc.add_heading('3.1 Structural Characteristics', level=2)
        
        table = self.doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        table.rows[0].cells[0].text = 'Metric'
        table.rows[0].cells[1].text = 'Value'
        
        # Data
        metrics_data = [
            ('Number of Nodes', str(orig_metrics['num_nodes'])),
            ('Number of Edges', str(orig_metrics['num_edges'])),
            ('Graph Density', f"{orig_metrics['density']:.4f}"),
            ('Average Degree', f"{orig_metrics['avg_degree']:.2f}"),
            ('Maximum In-Degree', str(orig_metrics['max_in_degree'])),
            ('Maximum Out-Degree', str(orig_metrics['max_out_degree']))
        ]
        
        for i, (metric, value) in enumerate(metrics_data, 1):
            table.rows[i].cells[0].text = metric
            table.rows[i].cells[1].text = value
        
        self.doc.add_paragraph()
        
        # Complexity analysis
        self.doc.add_heading('3.2 Complexity Analysis', level=2)
        
        complexity_text = f"The original graph demonstrates a redundancy ratio of {orig_metrics['redundancy_ratio']:.4f}, indicating that {orig_metrics['redundancy_ratio']*100:.1f}% of edges are transitive and potentially removable. The cyclomatic complexity is {orig_metrics['cyclomatic_complexity']}, and the degree entropy is {orig_metrics['degree_entropy']:.4f}, suggesting {'moderate' if orig_metrics['degree_entropy'] < 3 else 'high'} structural diversity."
        
        p = self.doc.add_paragraph(complexity_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_page_break()
    
    def _add_optimization_process(self):
        """Add optimization process description"""
        self.doc.add_heading('4. Optimization Process', level=1)
        
        process_text = "The optimization process was executed in two sequential phases to maximize graph reduction while ensuring correctness."
        p = self.doc.add_paragraph(process_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
        
        # Phase 1
        self.doc.add_heading('4.1 Phase 1: Transitive Reduction', level=2)
        phase1_text = "In the first phase, transitive reduction was applied using a modified Floyd-Warshall algorithm to identify and remove all transitive edges. This ensures that for any path of length > 1 from node u to node v, the direct edge (u,v) is removed if it exists. Edge attributes were preserved for all remaining edges."
        p = self.doc.add_paragraph(phase1_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
        
        # Phase 2
        self.doc.add_heading('4.2 Phase 2: Node Equivalence Merging', level=2)
        phase2_text = "In the second phase, nodes with identical signatures (predecessor and successor sets) were identified and merged. The merged node inherits the union of all edge attributes from its constituent nodes. This phase further reduces graph size without affecting semantic meaning."
        p = self.doc.add_paragraph(phase2_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_page_break()
    
    def _add_results_analysis(self, orig_metrics: Dict, opt_metrics: Dict):
        """Add results and analysis"""
        self.doc.add_heading('5. Results and Analysis', level=1)
        
        node_reduction = ((orig_metrics['num_nodes'] - opt_metrics['num_nodes']) / orig_metrics['num_nodes'] * 100)
        edge_reduction = ((orig_metrics['num_edges'] - opt_metrics['num_edges']) / orig_metrics['num_edges'] * 100)
        
        results_text = f"The optimization process successfully reduced the graph from {orig_metrics['num_nodes']} nodes and {orig_metrics['num_edges']} edges to {opt_metrics['num_nodes']} nodes and {opt_metrics['num_edges']} edges, representing reductions of {node_reduction:.1f}% and {edge_reduction:.1f}% respectively."
        
        p = self.doc.add_paragraph(results_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
        
        # Summary table
        self.doc.add_heading('5.1 Quantitative Results', level=2)
        
        table = self.doc.add_table(rows=4, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        headers = ['Metric', 'Original', 'Optimized', 'Change']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Data rows
        data = [
            ('Nodes', orig_metrics['num_nodes'], opt_metrics['num_nodes'], f"{node_reduction:.1f}%"),
            ('Edges', orig_metrics['num_edges'], opt_metrics['num_edges'], f"{edge_reduction:.1f}%"),
            ('Density', f"{orig_metrics['density']:.4f}", f"{opt_metrics['density']:.4f}", 
             f"{((opt_metrics['density'] - orig_metrics['density']) / orig_metrics['density'] * 100):.1f}%")
        ]
        
        for i, (metric, orig, opt, change) in enumerate(data, 1):
            table.rows[i].cells[0].text = metric
            table.rows[i].cells[1].text = str(orig)
            table.rows[i].cells[2].text = str(opt)
            table.rows[i].cells[3].text = change
        
        self.doc.add_page_break()
    
    def _add_detailed_metrics_comparison(self, orig_metrics: Dict, opt_metrics: Dict):
        """Add detailed metrics comparison table"""
        self.doc.add_heading('6. Detailed Metrics Comparison', level=1)
        
        # Create comprehensive comparison table
        table = self.doc.add_table(rows=15, cols=3)
        table.style = 'Medium Grid 3 Accent 1'
        
        # Headers
        headers = ['Metric', 'Original', 'Optimized']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Metrics data
        metrics = [
            ('Nodes', orig_metrics['num_nodes'], opt_metrics['num_nodes']),
            ('Edges', orig_metrics['num_edges'], opt_metrics['num_edges']),
            ('Leaf Nodes', orig_metrics['num_leaf_nodes'], opt_metrics['num_leaf_nodes']),
            ('Graph Density', f"{orig_metrics['density']:.4f}", f"{opt_metrics['density']:.4f}"),
            ('Avg Degree', f"{orig_metrics['avg_degree']:.2f}", f"{opt_metrics['avg_degree']:.2f}"),
            ('Max In-Degree', orig_metrics['max_in_degree'], opt_metrics['max_in_degree']),
            ('Max Out-Degree', orig_metrics['max_out_degree'], opt_metrics['max_out_degree']),
            ('Avg Path Length', f"{orig_metrics['avg_path_length']:.2f}", f"{opt_metrics['avg_path_length']:.2f}"),
            ('Diameter', orig_metrics['diameter'], opt_metrics['diameter']),
            ('Redundancy Ratio', f"{orig_metrics['redundancy_ratio']:.4f}", f"{opt_metrics['redundancy_ratio']:.4f}"),
            ('Efficiency Score', f"{orig_metrics['efficiency_score']:.4f}", f"{opt_metrics['efficiency_score']:.4f}"),
            ('Compactness Score', f"{orig_metrics['compactness_score']:.4f}", f"{opt_metrics['compactness_score']:.4f}"),
            ('Topological Complexity', orig_metrics['topological_complexity'], opt_metrics['topological_complexity']),
            ('Degree Entropy', f"{orig_metrics['degree_entropy']:.4f}", f"{opt_metrics['degree_entropy']:.4f}")
        ]
        
        for i, (metric, orig, opt) in enumerate(metrics, 1):
            table.rows[i].cells[0].text = metric
            table.rows[i].cells[1].text = str(orig)
            table.rows[i].cells[2].text = str(opt)
        
        self.doc.add_page_break()
    
    def _add_efficiency_analysis(self, orig_metrics: Dict, opt_metrics: Dict):
        """Add efficiency analysis section"""
        self.doc.add_heading('7. Efficiency Analysis', level=1)
        
        efficiency_gain = ((opt_metrics['efficiency_score'] - orig_metrics['efficiency_score']) * 100)
        redundancy_reduction = ((orig_metrics['redundancy_ratio'] - opt_metrics['redundancy_ratio']) * 100)
        
        efficiency_text = f"The optimization achieved a {efficiency_gain:.1f}% improvement in the composite efficiency score, rising from {orig_metrics['efficiency_score']:.4f} to {opt_metrics['efficiency_score']:.4f}. This improvement is primarily attributed to the {redundancy_reduction:.1f}% reduction in redundancy ratio, demonstrating effective removal of transitive edges."
        
        p = self.doc.add_paragraph(efficiency_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
        
        # Efficiency components
        self.doc.add_heading('7.1 Efficiency Score Components', level=2)
        
        formula_text = "The efficiency score is computed as a composite metric:\n\nE = [(1 - R) + (1 - D) + C] / 3\n\nWhere:\n• R = Redundancy Ratio\n• D = Graph Density\n• C = Compactness Score"
        
        p = self.doc.add_paragraph(formula_text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        self.doc.add_paragraph()
        
        # Component analysis
        component_text = f"Breaking down the efficiency components: Redundancy was reduced from {orig_metrics['redundancy_ratio']:.4f} to {opt_metrics['redundancy_ratio']:.4f}, density changed from {orig_metrics['density']:.4f} to {opt_metrics['density']:.4f}, and compactness improved from {orig_metrics['compactness_score']:.4f} to {opt_metrics['compactness_score']:.4f}. These improvements collectively contribute to the overall efficiency gain."
        
        p = self.doc.add_paragraph(component_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_page_break()
    
    def _add_critical_path_analysis(self, orig_metrics: Dict, opt_metrics: Dict):
        """Add critical path and bottleneck analysis"""
        self.doc.add_heading('8. Critical Path and Bottleneck Analysis', level=1)
        
        # Critical path
        self.doc.add_heading('8.1 Critical Path', level=2)
        
        cp_text = f"The critical path (longest path through the DAG) consists of {len(opt_metrics['critical_path'])} nodes in the optimized graph. This path represents the minimum time required for end-to-end processing in a parallel execution model."
        
        p = self.doc.add_paragraph(cp_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
        
        # Display critical path nodes
        if opt_metrics['critical_path']:
            cp_display = self.doc.add_paragraph()
            cp_display.add_run('Critical Path Nodes: ').bold = True
            cp_display.add_run(' → '.join(opt_metrics['critical_path'][:15]))
            if len(opt_metrics['critical_path']) > 15:
                cp_display.add_run(f' ... (and {len(opt_metrics["critical_path"]) - 15} more)')
        
        self.doc.add_paragraph()
        
        # Bottlenecks
        self.doc.add_heading('8.2 Bottleneck Nodes', level=2)
        
        bottleneck_text = f"Bottleneck analysis identified {len(opt_metrics['bottleneck_nodes'])} key nodes with high betweenness centrality. These nodes are critical for graph connectivity and represent potential performance bottlenecks in execution."
        
        p = self.doc.add_paragraph(bottleneck_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()
        
        # Display bottleneck nodes
        if opt_metrics['bottleneck_nodes']:
            bn_display = self.doc.add_paragraph()
            bn_display.add_run('Top Bottleneck Nodes: ').bold = True
            bn_display.add_run(', '.join(opt_metrics['bottleneck_nodes']))
        
        self.doc.add_page_break()
    
    def _add_conclusions(self, orig_metrics: Dict, opt_metrics: Dict):
        """Add conclusions section"""
        self.doc.add_heading('9. Conclusions', level=1)
        
        edge_reduction = ((orig_metrics['num_edges'] - opt_metrics['num_edges']) / orig_metrics['num_edges'] * 100)
        efficiency_gain = ((opt_metrics['efficiency_score'] - orig_metrics['efficiency_score']) * 100)
        
        conclusions = [
            f"This study successfully demonstrated the application of graph optimization techniques to reduce DAG complexity by {edge_reduction:.1f}% while preserving semantic relationships. The optimization process achieved a {efficiency_gain:.1f}% improvement in efficiency score, validating the effectiveness of combining transitive reduction with node equivalence merging.",
            
            "The results indicate that real-world DAGs often contain significant structural redundancy that can be eliminated without loss of information. The redundancy ratio decreased substantially, suggesting that many edges in the original graph were transitively implied by other paths.",
            
            f"Critical path analysis revealed that the optimized graph maintains {len(opt_metrics['critical_path'])} essential nodes along the longest execution path. Bottleneck identification highlighted {len(opt_metrics['bottleneck_nodes'])} key nodes that should be prioritized for performance optimization in practical applications.",
            
            "These findings have significant implications for large-scale graph processing applications, where reduced edge counts directly translate to lower memory requirements and faster traversal times. The methodology presented here provides a systematic approach to DAG optimization that can be applied across various domains."
        ]
        
        for text in conclusions:
            p = self.doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self.doc.add_paragraph()  # Spacing
        
        # Future work
        self.doc.add_heading('9.1 Future Work', level=2)
        
        future_text = "Future research directions include: (1) investigating the impact of optimization on query performance in graph databases, (2) developing heuristics for partial optimization in very large graphs, (3) analyzing the trade-offs between optimization depth and processing time, and (4) extending these techniques to dynamic graphs with temporal changes."
        
        p = self.doc.add_paragraph(future_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_page_break()
    
    def _add_references(self):
        """Add references section"""
        self.doc.add_heading('10. References', level=1)
        
        references = [
            "Aho, A. V., Garey, M. R., & Ullman, J. D. (1972). The transitive reduction of a directed graph. SIAM Journal on Computing, 1(2), 131-137.",
            
            "Cormen, T. H., Leiserson, C. E., Rivest, R. L., & Stein, C. (2009). Introduction to Algorithms (3rd ed.). MIT Press.",
            
            "Freeman, L. C. (1977). A set of measures of centrality based on betweenness. Sociometry, 40(1), 35-41.",
            
            "Hagberg, A., Swart, P., & S Chult, D. (2008). Exploring network structure, dynamics, and function using NetworkX. Los Alamos National Lab.(LANL), Los Alamos, NM (United States).",
            
            "Kahn, A. B. (1962). Topological sorting of large networks. Communications of the ACM, 5(11), 558-562.",
            
            "Mowshowitz, A. (1968). Entropy and the complexity of graphs: I. An index of the relative complexity of a graph. The Bulletin of Mathematical Biophysics, 30(1), 175-204.",
            
            "Tarjan, R. E. (1972). Depth-first search and linear graph algorithms. SIAM Journal on Computing, 1(2), 146-160.",
        ]
        
        for i, ref in enumerate(references, 1):
            p = self.doc.add_paragraph(f"[{i}] {ref}")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if i < len(references):
                self.doc.add_paragraph()  # Spacing between references

